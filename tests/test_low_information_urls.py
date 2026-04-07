from __future__ import annotations

from pathlib import Path
import tempfile
import unittest
from datetime import datetime

from docx import Document

from whatsapp_zip_to_docx.docx_writer import _render_url, _write_message_body
from whatsapp_zip_to_docx.parser import Message
from whatsapp_zip_to_docx.url_tools import UrlInfo


class LowInformationUrlTests(unittest.TestCase):
    def test_zoom_subdomain_is_classified_as_meeting(self) -> None:
        info = UrlInfo(
            original_url="https://us02web.zoom.us/j/123456789",
            final_url="https://us02web.zoom.us/j/123456789",
            domain="us02web.zoom.us",
            content_type="text/html",
            status=200,
        )

        self.assertEqual(info.kind, "meeting")
        self.assertEqual(info.source_label, "Zoom")

    def test_webex_meeting_path_is_classified_as_meeting(self) -> None:
        info = UrlInfo(
            original_url="https://acme.webex.com/meet/alain",
            final_url="https://acme.webex.com/meet/alain",
            domain="acme.webex.com",
            content_type="text/html",
            status=200,
        )

        self.assertEqual(info.kind, "meeting")
        self.assertEqual(info.source_label, "Webex")

    def test_youtube_short_link_is_classified_as_youtube(self) -> None:
        info = UrlInfo(
            original_url="https://youtu.be/B2J-uKbs40k",
            final_url="https://www.youtube.com/watch?v=B2J-uKbs40k",
            domain="www.youtube.com",
            content_type="text/html",
            status=200,
        )

        self.assertEqual(info.kind, "youtube")
        self.assertEqual(info.source_label, "YouTube")
        self.assertEqual(info.youtube_resource_type, "video")

    def test_dubb_link_is_classified_as_dubb(self) -> None:
        info = UrlInfo(
            original_url="https://iutum.dubb.com/v/recruiter_powerdale",
            final_url="https://iutum.dubb.com/v/recruiter_powerdale",
            domain="iutum.dubb.com",
            content_type="text/html",
            status=200,
            og_title="Powerdale Recruiter",
            og_description="Video from Dominique Dejonghe",
            og_site_name="Dubb",
        )

        self.assertEqual(info.kind, "dubb")
        self.assertEqual(info.source_label, "Dubb")
        self.assertEqual(info.dubb_title, "Powerdale Recruiter")
        self.assertEqual(info.dubb_creator, "Dominique Dejonghe")

    def test_dropbox_video_title_and_type_are_derived(self) -> None:
        info = UrlInfo(
            original_url="https://www.dropbox.com/s/ny1aefgho761z1o/Theodor%20Currentzis%20and%20his%20MusicAeterna%20in%20Konzerthaus.mp4?dl=0",
            final_url="https://www.dropbox.com/scl/fi/tmaokrcvqjn30917c7bow/Theodor-Currentzis-and-his-MusicAeterna-in-Konzerthaus.mp4?rlkey=x&dl=0",
            domain="www.dropbox.com",
            content_type="text/html",
            status=200,
            og_title="Dropbox",
        )

        self.assertEqual(info.kind, "dropbox")
        self.assertEqual(info.shared_title, "Theodor Currentzis and his MusicAeterna in Konzerthaus")
        self.assertEqual(info.shared_type_label, "Video partagee (Dropbox)")
        self.assertEqual(
            info.shared_video_source_url,
            "https://www.dropbox.com/scl/fi/tmaokrcvqjn30917c7bow/Theodor-Currentzis-and-his-MusicAeterna-in-Konzerthaus.mp4?rlkey=x&raw=1",
        )

    def test_icloud_numbers_title_type_and_owner_are_derived(self) -> None:
        info = UrlInfo(
            original_url="https://www.icloud.com/numbers/0h_GL9rJxw5KtxJtBncFtjD_Q#Pappie_Beethoven_Vioolconcerti",
            final_url="https://www.icloud.com/numbers/0h_GL9rJxw5KtxJtBncFtjD_Q#Pappie_Beethoven_Vioolconcerti",
            domain="www.icloud.com",
            content_type="text/html",
            status=200,
            og_title="Pappie Beethoven Vioolconcerti",
            og_description="Shared by Dominique Dejonghe",
            og_site_name="Numbers",
        )

        self.assertEqual(info.kind, "icloud")
        self.assertEqual(info.shared_title, "Pappie Beethoven Vioolconcerti")
        self.assertEqual(info.shared_type_label, "Document partage (iCloud Numbers)")
        self.assertEqual(info.shared_by, "Dominique Dejonghe")

    def test_google_docs_link_keeps_service_kind_even_with_error(self) -> None:
        info = UrlInfo(
            original_url="https://docs.google.com/document/d/abc/edit?usp=sharing",
            final_url="https://docs.google.com/document/d/abc/edit?usp=sharing",
            domain="docs.google.com",
            content_type="text/html",
            status=None,
            error="timed out",
        )

        self.assertEqual(info.kind, "google_drive")
        self.assertEqual(info.source_label, "Google Docs")
        self.assertEqual(info.shared_type_label, "Document partage (Google Docs)")

    def test_google_drive_video_file_uses_video_label(self) -> None:
        info = UrlInfo(
            original_url="https://drive.google.com/file/d/abc/view?usp=sharing",
            final_url="https://drive.google.com/file/d/abc/view?usp=sharing",
            domain="drive.google.com",
            content_type="text/html",
            status=200,
            og_title="Concert capture",
            og_type="video.other",
        )

        self.assertEqual(info.kind, "google_drive")
        self.assertEqual(info.shared_title, "Concert capture")
        self.assertEqual(info.shared_type_label, "Video partagee (Google Drive)")

    def test_linkedin_summary_is_capped_and_typed(self) -> None:
        description = " ".join(["How to make AI sound exactly like you forever."] * 12)
        info = UrlInfo(
            original_url="https://www.linkedin.com/posts/ruben-hassid_example",
            final_url="https://www.linkedin.com/posts/ruben-hassid_example",
            domain="www.linkedin.com",
            content_type="text/html",
            status=200,
            og_title="How to make AI sound exactly like you forever | Ruben Hassid",
            og_description=description,
            og_type="article",
            author="Ruben Hassid",
        )

        self.assertEqual(info.kind, "linkedin")
        self.assertEqual(info.source_label, "LinkedIn")
        self.assertEqual(info.linkedin_content_label, "Post LinkedIn")
        self.assertEqual(info.linkedin_source, "Ruben Hassid")
        self.assertLessEqual(len(info.linkedin_summary or ""), 220)
        self.assertTrue((info.linkedin_summary or "").endswith("…"))

    def test_x_status_gets_human_fallback_title(self) -> None:
        info = UrlInfo(
            original_url="https://x.com/Teslalightshows/status/1898012865342136664",
            final_url="https://x.com/Teslalightshows/status/1898012865342136664",
            domain="x.com",
            content_type="text/html",
            status=200,
        )

        self.assertEqual(info.kind, "x")
        self.assertEqual(info.source_label, "X")
        self.assertEqual(info.x_content_label, "Post X")
        self.assertEqual(info.x_source, "@Teslalightshows")
        self.assertEqual(info.x_title, "Post X de @Teslalightshows")

    def test_swr_video_uses_human_video_label(self) -> None:
        info = UrlInfo(
            original_url="https://www.swr.de/swrclassic/example.html",
            final_url="https://www.swr.de/swr-symphonieorchester/example.html",
            domain="www.swr.de",
            content_type="text/html",
            status=200,
            og_title="Currentzis LAB zu Mahlers Sinfonie Nr. 1",
            og_description="Mit Teodor Currentzis und Ilya Gaisin.",
            og_type="video.other",
            og_site_name="SWR",
        )

        self.assertEqual(info.kind, "swr")
        self.assertEqual(info.source_label, "SWR")
        self.assertEqual(info.swr_title, "Currentzis LAB zu Mahlers Sinfonie Nr. 1")
        self.assertEqual(info.swr_content_label, "Video SWR")

    def test_generic_webpage_type_labels_are_human(self) -> None:
        info = UrlInfo(
            original_url="https://example.com/video",
            final_url="https://example.com/video",
            domain="example.com",
            content_type="text/html",
            status=200,
            og_type="video.other",
        )

        self.assertEqual(info.kind, "webpage")
        self.assertEqual(info.webpage_content_type, "Video")

    def test_calendly_keeps_specific_title(self) -> None:
        info = UrlInfo(
            original_url="https://calendly.com/alain/coffee-chat",
            final_url="https://calendly.com/alain/coffee-chat",
            domain="calendly.com",
            content_type="text/html",
            status=200,
            og_title="Coffee chat with Alain - Calendly",
        )

        self.assertEqual(info.kind, "meeting")
        self.assertEqual(info.low_information_title, "Coffee chat with Alain")

    def test_generic_teams_title_is_suppressed(self) -> None:
        info = UrlInfo(
            original_url="https://teams.microsoft.com/l/meetup-join/abc",
            final_url="https://teams.microsoft.com/l/meetup-join/abc",
            domain="teams.microsoft.com",
            content_type="text/html",
            status=200,
            page_title="Join Microsoft Teams Meeting",
        )

        self.assertEqual(info.kind, "meeting")
        self.assertIsNone(info.low_information_title)

    def test_meeting_render_stays_compact_and_shows_error(self) -> None:
        url = "https://teams.microsoft.com/l/meetup-join/abc"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="teams.microsoft.com",
            content_type="text/html",
            status=None,
            page_title="Join Microsoft Teams Meeting",
            og_image="https://example.com/preview.png",
            error="timed out",
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            rendered = _render_url(document, url, {url: info}, Path(temp_dir))

        self.assertTrue(rendered)
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs],
            ["Teams", "Erreur d'acces: timed out"],
        )
        self.assertEqual(len(document.inline_shapes), 0)

    def test_inline_meeting_url_is_removed_when_metadata_block_is_rendered(self) -> None:
        url = "https://meet.google.com/qux-fkwe-xwe"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="meet.google.com",
            content_type="text/html",
            status=200,
        )
        message = Message(
            timestamp=datetime(2026, 2, 20, 15, 54, 0),
            author="Dominique",
            body=f"Breakout room {url}",
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[15:54] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["[15:54] D:", "Breakout room", "Google Meet"],
        )

    def test_inline_webpage_url_stays_in_message_text(self) -> None:
        url = "https://example.com/article"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="example.com",
            content_type="text/html",
            status=200,
            og_title="Useful article",
            og_description="Helpful context",
        )
        message = Message(
            timestamp=datetime(2026, 2, 20, 15, 54, 0),
            author="Dominique",
            body=f"Read this {url}",
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[15:54] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertIn(url, [paragraph.text for paragraph in document.paragraphs if paragraph.text][1])

    def test_inline_youtube_url_is_removed_when_metadata_block_is_rendered(self) -> None:
        url = "https://www.youtube.com/watch?v=huTSPanUlQM"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="www.youtube.com",
            content_type="text/html",
            status=200,
            og_title="The Art of Product Management with Sachin Rekhi (ENG'05 W'05)",
            og_site_name="YouTube",
            author="Wharton School",
        )
        message = Message(
            timestamp=datetime(2020, 8, 10, 17, 34, 25),
            author="Dominique",
            body=f"Interessante talk over Product Management: {url}",
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[17:34] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            [
                "[17:34] D:",
                "Interessante talk over Product Management:",
                "The Art of Product Management with Sachin Rekhi (ENG'05 W'05)",
                "Chaine (YouTube): Wharton School",
            ],
        )

    def test_unavailable_youtube_link_gets_plain_fallback_text(self) -> None:
        url = "https://youtu.be/B2J-uKbs40k"
        info = UrlInfo(
            original_url=url,
            final_url="https://www.youtube.com/watch?v=B2J-uKbs40k",
            domain="www.youtube.com",
            content_type="text/html",
            status=200,
        )
        message = Message(
            timestamp=datetime(2020, 9, 23, 14, 51, 22),
            author="Alain",
            body=url,
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[14:51] A: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Alain": "A"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["[14:51] A:", "YouTube", "This video isn't available any more"],
        )

    def test_inline_dubb_url_is_removed_when_metadata_block_is_rendered(self) -> None:
        url = "https://iutum.dubb.com/v/recruiter_powerdale"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="iutum.dubb.com",
            content_type="text/html",
            status=200,
            og_title="Powerdale Recruiter",
            og_description="Video from Dominique Dejonghe",
            og_site_name="Dubb",
        )
        message = Message(
            timestamp=datetime(2021, 4, 4, 8, 37, 59),
            author="Dominique",
            body=url,
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[08:37] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["[08:37] D:", "Powerdale Recruiter"],
        )

    def test_inline_dropbox_url_is_removed_when_metadata_block_is_rendered(self) -> None:
        url = "https://www.dropbox.com/s/zgnb749ndml7glz/IMG_7691.MOV?dl=0"
        info = UrlInfo(
            original_url=url,
            final_url=url,
            domain="www.dropbox.com",
            content_type="text/html",
            status=200,
            og_title="Dropbox",
        )
        message = Message(
            timestamp=datetime(2021, 6, 12, 14, 37, 21),
            author="Dominique",
            body=f"just for the we: {url}",
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[14:37] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["[14:37] D:", "just for the we:", "IMG 7691", "Video partagee (Dropbox)"],
        )

    def test_inline_swr_url_is_removed_when_metadata_block_is_rendered(self) -> None:
        url = "https://www.swr.de/swrclassic/example.html"
        info = UrlInfo(
            original_url=url,
            final_url="https://www.swr.de/swr-symphonieorchester/example.html",
            domain="www.swr.de",
            content_type="text/html",
            status=200,
            og_title="Currentzis LAB zu Mahlers Sinfonie Nr. 1",
            og_type="video.other",
            og_site_name="SWR",
        )
        message = Message(
            timestamp=datetime(2020, 2, 11, 20, 0, 0),
            author="Dominique",
            body=f"Concert: {url}",
            urls=[url],
        )
        document = Document()

        with tempfile.TemporaryDirectory() as temp_dir:
            _write_message_body(
                document,
                "[20:00] D: ",
                message.body,
                {url: info},
                spotify_mode="poeme",
                next_is_new_day=False,
                preview_dir=Path(temp_dir),
                messages=[message],
                initials_by_author={"Dominique": "D"},
            )

        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            ["[20:00] D:", "Concert:", "Currentzis LAB zu Mahlers Sinfonie Nr. 1", "Video SWR"],
        )


if __name__ == "__main__":
    unittest.main()
