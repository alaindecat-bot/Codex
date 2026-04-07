from __future__ import annotations

from datetime import date
from pathlib import Path
import re
import subprocess
import tempfile
from time import perf_counter

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .audio_transcription import AudioTranscript, is_supported_audio
from .parser import Message
from .perf import PerformanceRecorder
from .reply_analysis import ReplyCandidate
from .url_tools import UrlInfo, download_og_image

INLINE_URL_RE = re.compile(r"https?://\S+")

FRENCH_MONTHS = {
    1: "Janvier",
    2: "Fevrier",
    3: "Mars",
    4: "Avril",
    5: "Mai",
    6: "Juin",
    7: "Juillet",
    8: "Aout",
    9: "Septembre",
    10: "Octobre",
    11: "Novembre",
    12: "Decembre",
}

FRENCH_WEEKDAYS = {
    0: "Lundi",
    1: "Mardi",
    2: "Mercredi",
    3: "Jeudi",
    4: "Vendredi",
    5: "Samedi",
    6: "Dimanche",
}


def write_docx(
    output_path: Path,
    messages: list[Message],
    initials_by_author: dict[str, str],
    url_infos: dict[str, UrlInfo] | None = None,
    summary_lines: list[str] | None = None,
    attachment_links: dict[Path, str] | None = None,
    audio_transcripts: dict[Path, AudioTranscript] | None = None,
    reply_links: dict[int, ReplyCandidate] | None = None,
    spotify_mode: str = "simple",
    video_mode: str = "none",
    video_folder_link: str | None = None,
    body_font_size_pt: float = 10.5,
    column_count: int = 1,
    spotify_poem_columns: int = 1,
    spotify_poem_font_size_pt: float | None = None,
    profiler: PerformanceRecorder | None = None,
) -> None:
    document = Document()
    _configure_document_styles(document, body_font_size_pt=body_font_size_pt)
    _configure_document_layout(document, column_count=column_count)
    url_infos = url_infos or {}
    attachment_links = attachment_links or {}
    audio_transcripts = audio_transcripts or {}
    reply_links = reply_links or {}

    if summary_lines:
        _add_document_summary(document, summary_lines)

    with tempfile.TemporaryDirectory(prefix="whatsapp_zip_to_docx_previews_") as temp_dir_name:
        preview_dir = Path(temp_dir_name)
        current_month: tuple[int, int] | None = None
        current_day: date | None = None

        for index, message in enumerate(messages):
            next_message = messages[index + 1] if index + 1 < len(messages) else None
            message_month = (message.timestamp.year, message.timestamp.month)
            message_day = message.timestamp.date()

            if message_month != current_month:
                _add_month_heading(document, message.timestamp.year, message.timestamp.month)
                current_month = message_month
                current_day = None

            if message_day != current_day:
                _add_day_heading(document, message_day)
                current_day = message_day

            initial = initials_by_author.get(message.author, message.author[:1].upper())
            header = f"[{message.timestamp:%H:%M}] {initial}: "
            next_is_new_day = next_message is None or next_message.timestamp.date() != message_day
            reply_candidate = reply_links.get(index)

            if message.attachment and _is_visual_attachment(message.attachment.path):
                header_paragraph = _add_plain_paragraph(document, header)
                header_paragraph.paragraph_format.keep_with_next = True
                if reply_candidate is not None:
                    _add_reply_marker(document, messages, initials_by_author, reply_candidate)
                prepared = _prepare_visual_attachment(message.attachment.path, preview_dir)
                if prepared is not None:
                    _add_centered_picture(
                        document,
                        prepared,
                        height_cm=_attachment_height_cm(message.attachment.path),
                    )
                else:
                    _add_plain_paragraph(document, message.attachment.filename)
            elif message.attachment and _is_supported_video(message.attachment.path):
                header_paragraph = _add_plain_paragraph(document, header)
                header_paragraph.paragraph_format.keep_with_next = True
                if reply_candidate is not None:
                    _add_reply_marker(document, messages, initials_by_author, reply_candidate)
                preview = _render_video_preview(message.attachment.path, preview_dir)
                video_link = attachment_links.get(message.attachment.path)
                fallback_video_link = video_folder_link if video_mode == "drive" else None
                duration_label = _read_video_duration_label(message.attachment.path)
                if preview is not None:
                    _add_centered_picture(
                        document,
                        preview,
                        hyperlink=video_link or fallback_video_link,
                        blue_border=bool(video_link or fallback_video_link),
                    )
                else:
                    _add_plain_paragraph(document, message.attachment.filename)
            elif message.attachment and is_supported_audio(message.attachment.path):
                header_paragraph = _add_plain_paragraph(document, header)
                header_paragraph.paragraph_format.keep_with_next = True
                if reply_candidate is not None:
                    _add_reply_marker(document, messages, initials_by_author, reply_candidate)
                _add_audio_block(document, audio_transcripts.get(message.attachment.path))
            elif message.attachment:
                text = message.attachment.filename
                if reply_candidate is not None:
                    header_paragraph = _add_plain_paragraph(document, header)
                    header_paragraph.paragraph_format.keep_with_next = True
                    _add_reply_marker(document, messages, initials_by_author, reply_candidate)
                    _add_plain_paragraph(document, text)
                else:
                    _add_plain_paragraph(document, f"{header}{text}")
            else:
                _write_message_body(
                    document,
                    header,
                    message.body,
                    url_infos,
                    spotify_mode,
                    next_is_new_day,
                    preview_dir,
                    messages=messages,
                    initials_by_author=initials_by_author,
                    reply_candidate=reply_candidate,
                    spotify_poem_columns=spotify_poem_columns,
                    spotify_poem_font_size_pt=spotify_poem_font_size_pt,
                    profiler=profiler,
                )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _is_supported_image(path: Path) -> bool:
    return path.suffix.lower() in {".jpg", ".jpeg", ".png", ".gif", ".bmp"}


def _is_renderable_pdf(path: Path) -> bool:
    return path.suffix.lower() == ".pdf"


def _is_supported_video(path: Path) -> bool:
    return path.suffix.lower() in {".mp4", ".mov", ".m4v"}


def _is_visual_attachment(path: Path) -> bool:
    return _is_supported_image(path) or _is_renderable_pdf(path)


def _is_url_only_message(body: str) -> bool:
    stripped = body.strip()
    return stripped.startswith("http://") or stripped.startswith("https://")


def _append_url_metadata(
    document: Document,
    info: UrlInfo | None,
    spotify_mode: str = "simple",
    spotify_poem_columns: int = 1,
    spotify_poem_font_size_pt: float | None = None,
) -> bool:
    if info is None:
        return False

    if info.kind == "unreachable":
        _add_url_block_heading(document, info.source_label, hyperlink=info.final_url or info.original_url)
        if info.error:
            _add_url_detail_line(document, f"Erreur d'acces: {info.error}")
        return True

    if info.kind == "meeting":
        _add_url_block_heading(document, info.source_label, hyperlink=info.final_url or info.original_url)
        if info.error:
            _add_url_detail_line(document, f"Erreur d'acces: {info.error}")
        return True

    if info.kind == "spotify":
        details = []
        title = info.spotify_title
        if title:
            details.append(title)
        if info.spotify_artists:
            details.append(f"Artiste(s): {info.spotify_artists}")
        if spotify_mode != "poeme" or not info.spotify_lyrics:
            details.append(
                "Paroles trouvables: oui"
                if info.lyrics_searchable
                else "Paroles trouvables: non"
            )
        if not details:
            return False
        _add_url_title_line(document, details[0], hyperlink=info.final_url)
        for line in details[1:]:
            _add_url_detail_line(document, line)
        if spotify_mode == "poeme" and info.spotify_lyrics:
            _add_spotify_lyrics_poem(
                document,
                info.spotify_lyrics,
                column_count=spotify_poem_columns,
                font_size_pt=spotify_poem_font_size_pt,
            )
        return True

    if info.kind == "youtube":
        title = info.youtube_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, info.source_label, hyperlink=info.final_url or info.original_url)
        details = []
        if info.youtube_creator:
            details.append(f"Chaine (YouTube): {info.youtube_creator}")
        if not title and not info.youtube_creator:
            details.append("This video isn't available any more")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind == "dubb":
        title = info.dubb_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, info.source_label, hyperlink=info.final_url or info.original_url)
        details = []
        if not title and info.dubb_creator:
            details.append(f"Source (Dubb): {info.dubb_creator}")
        elif not title:
            details.append("It's a Dubb video link")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind in {"dropbox", "google_drive", "icloud"}:
        title = info.shared_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, info.source_label, hyperlink=info.final_url or info.original_url)
        details = []
        if info.shared_type_label:
            details.append(info.shared_type_label)
        if info.shared_by:
            details.append(f"Partage par: {info.shared_by}")
        if info.error:
            details.append(f"Erreur d'acces: {info.error}")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind == "linkedin":
        title = info.linkedin_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, "LinkedIn", hyperlink=info.final_url or info.original_url)
        details = []
        if info.linkedin_content_label:
            details.append(info.linkedin_content_label)
        if info.linkedin_source:
            details.append(f"Source: {info.linkedin_source}")
        if info.linkedin_summary:
            details.append(f"Resume: {info.linkedin_summary}")
        if info.error:
            details.append(f"Erreur d'acces: {info.error}")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind == "x":
        title = info.x_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, "X", hyperlink=info.final_url or info.original_url)
        details = []
        if info.x_content_label and info.x_content_label != title:
            details.append(info.x_content_label)
        if info.x_source and (title is None or info.x_source not in title):
            details.append(f"Source: {info.x_source}")
        if info.x_summary:
            details.append(f"Resume: {info.x_summary}")
        if info.error:
            details.append(f"Erreur d'acces: {info.error}")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind == "swr":
        title = info.swr_title
        if title:
            _add_url_title_line(document, title, hyperlink=info.final_url or info.original_url)
        else:
            _add_url_block_heading(document, "SWR", hyperlink=info.final_url or info.original_url)
        details = []
        if info.swr_content_label:
            details.append(info.swr_content_label)
        if info.swr_summary:
            details.append(f"Resume: {info.swr_summary}")
        if info.error:
            details.append(f"Erreur d'acces: {info.error}")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(title or details)

    if info.kind == "facebook":
        _add_url_block_heading(document, "Facebook", hyperlink=info.final_url)
        details = []
        if info.facebook_content_type:
            details.append(f"Type: {info.facebook_content_type}")
        if info.facebook_source:
            details.append(f"Source: {info.facebook_source}")
        if info.facebook_summary:
            details.append(f"Resume: {info.facebook_summary}")
        if info.can_embed_post_image is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(details)

    if info.kind == "webpage":
        heading = info.web_source or info.source_label or "Lien"
        _add_url_block_heading(document, heading, hyperlink=info.final_url)
        details = []
        if info.webpage_content_type:
            details.append(f"Type: {info.webpage_content_type}")
        if info.web_author:
            details.append(f"Auteur: {info.web_author}")
        if info.web_summary:
            details.append(f"Resume: {info.web_summary}")
        if info.og_image and info.image_fetchable is False:
            details.append("Image recuperable: non")
        for line in details:
            _add_url_detail_line(document, line)
        return bool(details)

    return False


def _write_message_body(
    document: Document,
    header: str,
    body: str,
    url_infos: dict[str, UrlInfo],
    spotify_mode: str,
    next_is_new_day: bool,
    preview_dir: Path,
    messages: list[Message],
    initials_by_author: dict[str, str],
    reply_candidate: ReplyCandidate | None = None,
    spotify_poem_columns: int = 1,
    spotify_poem_font_size_pt: float | None = None,
    profiler: PerformanceRecorder | None = None,
) -> None:
    lines = body.splitlines() or [body]
    if _is_poem(body):
        _write_poem_message(document, header.rstrip(), body, messages, initials_by_author, reply_candidate)
        return

    wrote_header = False
    appended_metadata = False

    for line in lines:
        stripped = line.strip()
        line_urls = INLINE_URL_RE.findall(line)
        inline_text = _strip_redundant_inline_urls(line, line_urls, url_infos)
        if not wrote_header:
            if reply_candidate is not None:
                header_paragraph = _add_plain_paragraph(document, header.rstrip())
                header_paragraph.paragraph_format.keep_with_next = True
                _add_reply_marker(document, messages, initials_by_author, reply_candidate)
                wrote_header = True
                if not line_urls:
                    _add_plain_paragraph(document, line)
                    continue
            if line_urls:
                if len(line_urls) == 1 and stripped == line_urls[0]:
                    _add_plain_paragraph(document, header.rstrip())
                    wrote_header = True
                else:
                    _add_plain_paragraph(document, header.rstrip())
                    if inline_text.strip():
                        _add_text_with_hyperlinks(document, inline_text)
                    wrote_header = True
            else:
                _add_plain_paragraph(document, f"{header}{line}")
                wrote_header = True
                continue

        if line_urls:
            if len(line_urls) == 1 and stripped == line_urls[0]:
                appended_metadata = _render_url(
                    document,
                    line_urls[0],
                    url_infos,
                    preview_dir,
                    spotify_mode=spotify_mode,
                    spotify_poem_columns=spotify_poem_columns,
                    spotify_poem_font_size_pt=spotify_poem_font_size_pt,
                    profiler=profiler,
                ) or appended_metadata
            else:
                for url in line_urls:
                    appended_metadata = _render_url(
                        document,
                        url,
                        url_infos,
                        preview_dir,
                        spotify_mode=spotify_mode,
                        spotify_poem_columns=spotify_poem_columns,
                        spotify_poem_font_size_pt=spotify_poem_font_size_pt,
                        render_fallback_quote=False,
                        profiler=profiler,
                    ) or appended_metadata
        elif stripped:
            _add_plain_paragraph(document, line)
        else:
            document.add_paragraph("")

    if not wrote_header:
        _add_plain_paragraph(document, header.rstrip())

    if appended_metadata and not next_is_new_day:
        document.add_paragraph("")


def _strip_redundant_inline_urls(line: str, line_urls: list[str], url_infos: dict[str, UrlInfo]) -> str:
    if not line_urls:
        return line
    parts: list[str] = []
    last_end = 0
    changed = False
    for match in INLINE_URL_RE.finditer(line):
        url = match.group(0)
        parts.append(line[last_end:match.start()])
        if _should_replace_inline_url(url, url_infos.get(url)):
            changed = True
        else:
            parts.append(url)
        last_end = match.end()
    parts.append(line[last_end:])
    if not changed:
        return line
    return _normalize_inline_text_after_url_removal("".join(parts))


def _should_replace_inline_url(url: str, info: UrlInfo | None) -> bool:
    if info is None:
        return False
    if info.kind == "meeting":
        return True
    if info.kind == "youtube":
        return bool(info.youtube_title or info.youtube_creator or info.og_image)
    if info.kind == "dubb":
        return bool(info.dubb_title or info.dubb_creator or info.og_image)
    if info.kind in {"dropbox", "google_drive", "icloud"}:
        return bool(info.shared_title or info.shared_type_label or info.og_image)
    if info.kind == "linkedin":
        return bool(info.linkedin_title or info.linkedin_content_label or info.og_image or info.error)
    if info.kind == "x":
        return bool(info.x_title or info.x_content_label or info.og_image or info.error)
    if info.kind == "swr":
        return bool(info.swr_title or info.swr_content_label or info.og_image or info.error)
    return False


def _normalize_inline_text_after_url_removal(text: str) -> str:
    compacted = re.sub(r"[ \t]{2,}", " ", text)
    compacted = re.sub(r" +([,.;:!?])", r"\1", compacted)
    return compacted.strip()


def _configure_document_styles(document: Document, body_font_size_pt: float = 10.5) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(body_font_size_pt)
    r_fonts = normal_style.element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:cs"), "Calibri")


def _configure_document_layout(document: Document, column_count: int = 1) -> None:
    if column_count <= 1:
        return
    section = document.sections[0]
    _set_section_columns(section, column_count)


def _set_section_columns(section, column_count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(column_count))
    cols.set(qn("w:space"), "709" if column_count > 1 else "720")


def _add_month_heading(document: Document, year: int, month: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{FRENCH_MONTHS[month]} {year}")
    _apply_run_format(run, bold=True, size=Pt(15), color=RGBColor(0x4A, 0x4A, 0x4A))


def _add_document_summary(document: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0 if index else 6)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        _apply_run_format(run, italic=True, size=Pt(10))
    document.add_paragraph("")


def _add_day_heading(document: Document, value: date) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{FRENCH_WEEKDAYS[value.weekday()]} {value.day}")
    _apply_run_format(run, bold=True, size=Pt(11), color=RGBColor(0x3F, 0x5B, 0x6B))


def _add_url_block_heading(document: Document, label: str, hyperlink: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_with_next = True
    if hyperlink:
        _add_hyperlink(paragraph, hyperlink, label, bold=True, courier=True)
    else:
        run = paragraph.add_run(label)
        _apply_run_format(run, bold=True, size=Pt(10.5), courier=True)


def _add_url_title_line(document: Document, text: str, hyperlink: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    if hyperlink:
        _add_hyperlink(paragraph, hyperlink, text, bold=True, courier=True)
    else:
        run = paragraph.add_run(text)
        _apply_run_format(run, bold=True, courier=True)


def _add_url_detail_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _apply_run_format(run, courier=True)


def _add_spotify_lyrics_poem(
    document: Document,
    lyrics: str,
    column_count: int = 1,
    font_size_pt: float | None = None,
) -> None:
    document.add_paragraph("")
    stanzas: list[list[str]] = []
    current: list[str] = []
    for raw_line in lyrics.splitlines():
        line = raw_line.strip()
        if line:
            current.append(line)
        else:
            if current:
                stanzas.append(current)
                current = []
    if current:
        stanzas.append(current)

    if column_count > 1:
        two_col_section = document.add_section(WD_SECTION.CONTINUOUS)
        _set_section_columns(two_col_section, column_count)

    stanza_break_index = _spotify_column_break_index(stanzas) if column_count > 1 else None
    for stanza_index, stanza in enumerate(stanzas):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.55)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(4)
        for index, line in enumerate(stanza):
            run = paragraph.add_run(line)
            _apply_run_format(run, italic=True, size=Pt(font_size_pt) if font_size_pt else None)
            if index < len(stanza) - 1:
                run.add_break()
        if stanza_break_index is not None and stanza_index == stanza_break_index:
            break_paragraph = document.add_paragraph()
            break_paragraph.paragraph_format.left_indent = Inches(0.55)
            break_run = break_paragraph.add_run()
            _apply_run_format(break_run, italic=True, size=Pt(font_size_pt) if font_size_pt else None)
            break_run.add_break(WD_BREAK.COLUMN)

    if column_count > 1:
        one_col_section = document.add_section(WD_SECTION.CONTINUOUS)
        _set_section_columns(one_col_section, 1)


def _spotify_column_break_index(stanzas: list[list[str]]) -> int | None:
    if len(stanzas) < 2:
        return None
    total_lines = sum(len(stanza) for stanza in stanzas)
    if total_lines < 8:
        return None
    running = 0
    best_index = None
    best_distance = None
    target = total_lines / 2
    for index, stanza in enumerate(stanzas[:-1]):
        running += len(stanza)
        distance = abs(target - running)
        if best_distance is None or distance < best_distance:
            best_distance = distance
            best_index = index
    return best_index


def _add_plain_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _apply_run_format(run)
    return paragraph


def _add_text_with_hyperlinks(document: Document, text: str):
    paragraph = document.add_paragraph()
    last_end = 0
    for match in INLINE_URL_RE.finditer(text):
        prefix = text[last_end:match.start()]
        if prefix:
            run = paragraph.add_run(prefix)
            _apply_run_format(run)
        url = match.group(0)
        _add_hyperlink(paragraph, url, url)
        last_end = match.end()
    suffix = text[last_end:]
    if suffix:
        run = paragraph.add_run(suffix)
        _apply_run_format(run)
    return paragraph


def _add_reply_marker(
    document: Document,
    messages: list[Message],
    initials_by_author: dict[str, str],
    candidate: ReplyCandidate,
) -> None:
    if candidate.intervening_count == 0:
        return
    prompt = messages[candidate.prompt_index]
    initial = initials_by_author.get(prompt.author, prompt.author[:1].upper())
    excerpt = _truncate_reply_excerpt(_single_line(prompt.body))
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"↳ Reponse a [{prompt.timestamp:%H:%M}] {initial}: {excerpt}")
    _apply_run_format(run, italic=True, size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))


def _add_quote_paragraph(document: Document, text: str, hyperlink: str | None = None) -> None:
    paragraph = document.add_paragraph(style="Intense Quote")
    if hyperlink:
        _add_hyperlink(paragraph, hyperlink, text)
    else:
        run = paragraph.add_run(text)
        _apply_run_format(run)


def _add_attachment_caption(document: Document, text: str, subtle: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(2 if not subtle else 0)
    paragraph.paragraph_format.space_after = Pt(6 if not subtle else 4)
    run = paragraph.add_run(text)
    _apply_run_format(
        run,
        italic=True,
        size=Pt(9 if not subtle else 8.5),
        color=RGBColor(0x66, 0x66, 0x66) if subtle else None,
    )


def _add_audio_block(document: Document, transcript: AudioTranscript | None) -> None:
    label_paragraph = document.add_paragraph()
    label_paragraph.paragraph_format.left_indent = Inches(0.25)
    label_paragraph.paragraph_format.space_before = Pt(1)
    label_paragraph.paragraph_format.space_after = Pt(1)
    label_paragraph.paragraph_format.keep_with_next = True
    label_run = label_paragraph.add_run("Audio")
    _apply_run_format(label_run, italic=True, size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    transcript_label = "Transcription"
    if transcript and transcript.language:
        transcript_label = f"Transcription ({transcript.language})"
    transcript_label_paragraph = document.add_paragraph()
    transcript_label_paragraph.paragraph_format.left_indent = Inches(0.35)
    transcript_label_paragraph.paragraph_format.space_before = Pt(0)
    transcript_label_paragraph.paragraph_format.space_after = Pt(1)
    transcript_label_paragraph.paragraph_format.keep_with_next = True
    transcript_label_run = transcript_label_paragraph.add_run(transcript_label)
    _apply_run_format(transcript_label_run, italic=True, size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    body_paragraph = document.add_paragraph()
    body_paragraph.paragraph_format.left_indent = Inches(0.45)
    body_paragraph.paragraph_format.space_before = Pt(0)
    body_paragraph.paragraph_format.space_after = Pt(6)
    body_text = "Transcription indisponible"
    if transcript and transcript.available and transcript.text:
        body_text = transcript.text
    body_run = body_paragraph.add_run(body_text)
    _apply_run_format(body_run, size=Pt(10))


def _append_web_preview(
    document: Document,
    info: UrlInfo,
    preview_dir: Path,
    hyperlink: str | None = None,
    blue_border: bool = False,
) -> None:
    _append_web_preview_with_profile(
        document,
        info,
        preview_dir,
        hyperlink=hyperlink,
        blue_border=blue_border,
        profiler=None,
    )


def _append_web_preview_with_profile(
    document: Document,
    info: UrlInfo,
    preview_dir: Path,
    hyperlink: str | None = None,
    blue_border: bool = False,
    profiler: PerformanceRecorder | None = None,
) -> None:
    started = None
    if profiler is not None:
        started = perf_counter()
    rendered = False
    preview_path = download_og_image(info, preview_dir / _safe_preview_dirname(info), timeout=6.0)
    if preview_path is not None:
        prepared = _prepare_image_for_render(preview_path, workspace=preview_dir)
        height_cm = 11 if info.kind == "facebook" and _is_portrait_image(prepared) else 7
        rendered = _add_centered_picture(
            document,
            prepared,
            hyperlink=hyperlink,
            blue_border=blue_border,
            height_cm=height_cm,
        )
    if profiler is not None and started is not None:
        profiler.record(
            "preview",
            "preview.web_image",
            perf_counter() - started,
            url=info.final_url or info.original_url,
            kind=info.kind,
            domain=info.domain,
            ok=rendered,
        )


def _append_remote_video_preview(
    document: Document,
    info: UrlInfo,
    preview_dir: Path,
    hyperlink: str | None = None,
    blue_border: bool = False,
    profiler: PerformanceRecorder | None = None,
) -> bool:
    started = None
    if profiler is not None:
        started = perf_counter()
    preview = _render_remote_video_preview(info, preview_dir)
    rendered = False
    if preview is not None:
        rendered = _add_centered_picture(document, preview, hyperlink=hyperlink, blue_border=blue_border, height_cm=7)
    if profiler is not None and started is not None:
        profiler.record(
            "preview",
            "preview.remote_video",
            perf_counter() - started,
            url=info.final_url or info.original_url,
            kind=info.kind,
            domain=info.domain,
            ok=rendered,
        )
    return rendered


def _attachment_height_cm(attachment_path: Path) -> int:
    if _is_renderable_pdf(attachment_path):
        return 14
    return 7


def _render_url(
    document: Document,
    url: str,
    url_infos: dict[str, UrlInfo],
    preview_dir: Path,
    spotify_mode: str = "simple",
    spotify_poem_columns: int = 1,
    spotify_poem_font_size_pt: float | None = None,
    render_fallback_quote: bool = True,
    profiler: PerformanceRecorder | None = None,
) -> bool:
    info = url_infos.get(url)
    if info is None:
        if render_fallback_quote:
            _add_quote_paragraph(document, url, hyperlink=url)
        return False

    if info.kind == "spotify":
        return _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )

    if info.kind == "facebook":
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        _append_web_preview_with_profile(
            document,
            info,
            preview_dir,
            hyperlink=info.final_url,
            blue_border=True,
            profiler=profiler,
        )
        return rendered or bool(info.og_image)

    if info.kind == "youtube":
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        if info.og_image:
            _append_web_preview_with_profile(
                document,
                info,
                preview_dir,
                hyperlink=info.final_url,
                blue_border=True,
                profiler=profiler,
            )
        return rendered or bool(info.og_image)

    if info.kind == "dubb":
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        if info.og_image:
            _append_web_preview_with_profile(
                document,
                info,
                preview_dir,
                hyperlink=info.final_url,
                blue_border=True,
                profiler=profiler,
            )
        return rendered or bool(info.og_image)

    if info.kind in {"dropbox", "google_drive", "icloud"}:
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        if info.og_image:
            _append_web_preview_with_profile(
                document,
                info,
                preview_dir,
                hyperlink=info.final_url,
                blue_border=True,
                profiler=profiler,
            )
        elif info.shared_video_source_url:
            _append_remote_video_preview(
                document,
                info,
                preview_dir,
                hyperlink=info.final_url,
                blue_border=True,
                profiler=profiler,
            )
        return rendered or bool(info.og_image)

    if info.kind in {"linkedin", "x", "swr"}:
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        if info.og_image:
            _append_web_preview_with_profile(
                document,
                info,
                preview_dir,
                hyperlink=info.final_url,
                blue_border=True,
                profiler=profiler,
            )
        return rendered or bool(info.og_image)

    if info.kind == "meeting":
        return _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )

    if info.kind == "webpage" and info.og_image:
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        _append_web_preview_with_profile(
            document,
            info,
            preview_dir,
            hyperlink=info.final_url,
            blue_border=True,
            profiler=profiler,
        )
        return rendered or True

    if info.kind == "webpage":
        rendered = _append_url_metadata(
            document,
            info,
            spotify_mode=spotify_mode,
            spotify_poem_columns=spotify_poem_columns,
            spotify_poem_font_size_pt=spotify_poem_font_size_pt,
        )
        if render_fallback_quote:
            _add_quote_paragraph(document, url, hyperlink=url)
        return rendered

    rendered = _append_url_metadata(
        document,
        info,
        spotify_mode=spotify_mode,
        spotify_poem_columns=spotify_poem_columns,
        spotify_poem_font_size_pt=spotify_poem_font_size_pt,
    )
    if render_fallback_quote and info.kind != "unreachable":
        _add_quote_paragraph(document, url, hyperlink=url)
    return rendered


def _add_centered_picture(
    document: Document,
    image_path: Path,
    hyperlink: str | None = None,
    blue_border: bool = False,
    height_cm: int = 7,
) -> bool:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.keep_together = True
    if blue_border:
        _apply_blue_paragraph_border(paragraph)
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), height=Cm(height_cm))
    except (FileNotFoundError, UnrecognizedImageError):
        paragraph._element.getparent().remove(paragraph._element)
        return False
    if hyperlink:
        _wrap_run_in_hyperlink(paragraph, run, hyperlink)
    return True


def _safe_preview_dirname(info: UrlInfo) -> Path:
    stem = re.sub(r"[^a-zA-Z0-9_-]+", "_", info.domain or "preview").strip("_") or "preview"
    return Path(stem)


def _is_poem(body: str) -> bool:
    if not body:
        return False
    lines = [line.strip() for line in body.splitlines()]
    non_empty = [line for line in lines if line]
    if len(non_empty) < 4:
        return False
    average_length = sum(len(line) for line in non_empty) / len(non_empty)
    short_line_count = sum(1 for line in non_empty if len(line) <= 55)
    return average_length <= 55 and short_line_count >= max(3, len(non_empty) - 1)


def _write_poem_message(
    document: Document,
    header: str,
    body: str,
    messages: list[Message],
    initials_by_author: dict[str, str],
    reply_candidate: ReplyCandidate | None = None,
) -> None:
    _add_plain_paragraph(document, header)
    if reply_candidate is not None:
        _add_reply_marker(document, messages, initials_by_author, reply_candidate)
    stanzas = []
    current: list[str] = []
    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if line.strip():
            current.append(line.strip())
        else:
            if current:
                stanzas.append(current)
                current = []
    if current:
        stanzas.append(current)

    for stanza in stanzas:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(6)
        for index, line in enumerate(stanza):
            run = paragraph.add_run(line)
            _apply_run_format(run, italic=True)
            if index < len(stanza) - 1:
                run.add_break()


def _prepare_visual_attachment(attachment_path: Path, workspace: Path) -> Path | None:
    if _is_supported_image(attachment_path):
        return _prepare_image_for_render(attachment_path, workspace=workspace)
    if _is_renderable_pdf(attachment_path):
        preview = _render_pdf_preview(attachment_path, workspace)
        if preview is None:
            return None
        return _prepare_image_for_render(preview, workspace=workspace)
    return None


def _prepare_image_for_render(image_path: Path, workspace: Path) -> Path:
    rotation = _image_rotation_degrees(image_path)
    prepared = image_path
    if rotation != 0:
        rotated_path = workspace / f"rotated_{rotation}_{image_path.name}"
        try:
            subprocess.run(
                ["/usr/bin/sips", "-r", str(rotation), str(image_path), "--out", str(rotated_path)],
                check=True,
                capture_output=True,
                text=True,
            )
        except (OSError, subprocess.CalledProcessError):
            prepared = image_path
        else:
            prepared = rotated_path if rotated_path.exists() else image_path
    return _ensure_docx_compatible_image(prepared, workspace)


def _render_pdf_preview(pdf_path: Path, workspace: Path) -> Path | None:
    output_dir = workspace / f"pdf_{pdf_path.stem}"
    output_dir.mkdir(parents=True, exist_ok=True)
    return _run_quicklook_thumbnail(pdf_path, output_dir)


def _render_video_preview(video_path: Path, workspace: Path) -> Path | None:
    output_dir = workspace / f"video_{video_path.stem}"
    output_dir.mkdir(parents=True, exist_ok=True)
    preview = _run_quicklook_thumbnail(video_path, output_dir)
    if preview is None:
        return None
    return _prepare_image_for_render(preview, workspace=workspace)


def _render_remote_video_preview(info: UrlInfo, workspace: Path) -> Path | None:
    source_url = info.shared_video_source_url
    if not source_url:
        return None
    output_dir = workspace / f"remote_video_{_safe_preview_dirname(info)}"
    output_dir.mkdir(parents=True, exist_ok=True)
    preview_path = output_dir / "preview.png"
    try:
        subprocess.run(
            [
                "/opt/homebrew/bin/ffmpeg",
                "-y",
                "-ss",
                "1",
                "-i",
                source_url,
                "-frames:v",
                "1",
                str(preview_path),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
    except (OSError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None
    if not preview_path.exists():
        return None
    return _prepare_image_for_render(preview_path, workspace=workspace)


def _read_video_duration_label(video_path: Path) -> str | None:
    try:
        result = subprocess.run(
            ["mdls", "-name", "kMDItemDurationSeconds", str(video_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    for line in result.stdout.splitlines():
        if "kMDItemDurationSeconds" not in line or "=" not in line:
            continue
        raw = line.split("=", 1)[1].strip()
        if raw == "(null)":
            return None
        try:
            total_seconds = int(round(float(raw)))
        except ValueError:
            return None
        minutes, seconds = divmod(total_seconds, 60)
        return f"{minutes}:{seconds:02d}"
    return None


def _run_quicklook_thumbnail(source_path: Path, output_dir: Path) -> Path | None:
    try:
        subprocess.run(
            ["qlmanage", "-t", "-s", "2000", "-o", str(output_dir), str(source_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None

    candidates = sorted(output_dir.glob(f"{source_path.name}.*")) + sorted(output_dir.glob(f"{source_path.stem}.*"))
    for candidate in candidates:
        if candidate.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            return candidate
    return None


def _image_rotation_degrees(image_path: Path) -> int:
    orientation = _read_image_orientation(image_path)
    if orientation == 3:
        return 180
    if orientation == 6:
        return 90
    if orientation == 8:
        return 270
    return 0


def _read_image_orientation(image_path: Path) -> int | None:
    try:
        result = subprocess.run(
            ["/usr/bin/sips", "-g", "orientation", str(image_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    for line in result.stdout.splitlines():
        line = line.strip()
        if line.startswith("orientation:"):
            value = line.split(":", 1)[1].strip()
            if value == "<nil>":
                return None
            try:
                return int(value)
            except ValueError:
                return None
    return None


def _ensure_docx_compatible_image(image_path: Path, workspace: Path) -> Path:
    if image_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp"}:
        return image_path
    converted_path = workspace / f"{image_path.stem}.png"
    try:
        subprocess.run(
            ["/usr/bin/sips", "-s", "format", "png", str(image_path), "--out", str(converted_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return image_path
    return converted_path if converted_path.exists() else image_path


def _read_image_dimensions(image_path: Path) -> tuple[int, int] | None:
    try:
        result = subprocess.run(
            ["/usr/bin/sips", "-g", "pixelWidth", "-g", "pixelHeight", str(image_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    width = None
    height = None
    for line in result.stdout.splitlines():
        line = line.strip()
        if line.startswith("pixelWidth:"):
            width = int(line.split(":")[1].strip())
        elif line.startswith("pixelHeight:"):
            height = int(line.split(":")[1].strip())
    if width is None or height is None:
        return None
    return width, height


def _is_portrait_image(image_path: Path) -> bool:
    dimensions = _read_image_dimensions(image_path)
    if dimensions is None:
        return False
    width, height = dimensions
    return height > width


def _apply_run_format(
    run,
    bold: bool = False,
    italic: bool = False,
    size: Pt | None = None,
    color: RGBColor | None = None,
    courier: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    font_name = "Courier New" if courier else "Calibri"
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def _add_hyperlink(
    paragraph,
    url: str,
    text: str,
    bold: bool = False,
    courier: bool = False,
) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(r_style)

    r_fonts = OxmlElement("w:rFonts")
    font_name = "Courier New" if courier else "Calibri"
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    run_properties.append(r_fonts)

    if bold:
        run_properties.append(OxmlElement("w:b"))

    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _wrap_run_in_hyperlink(paragraph, run, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    paragraph._p.remove(run._r)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def _apply_blue_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), "4F81BD")
        p_bdr.append(element)
    p_pr.append(p_bdr)


def _single_line(text: str) -> str:
    return " ".join(part.strip() for part in text.splitlines() if part.strip())


def _truncate_reply_excerpt(text: str, max_length: int = 70) -> str:
    if len(text) <= max_length:
        return text
    return text[: max_length - 1].rstrip() + "…"


def _build_video_caption(label: str, duration: str | None) -> str:
    if duration:
        return f"{label} - {duration}"
    return label
